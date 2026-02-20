"""
Build a properly formatted Word document for bioRxiv submission.
Uses python-docx for full control over styles, tables, and layout.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page layout ──────────────────────────────────────────────────────────────
section = doc.sections[0]
section.page_height = Inches(11)
section.page_width  = Inches(8.5)
section.left_margin   = Inches(1)
section.right_margin  = Inches(1)
section.top_margin    = Inches(1)
section.bottom_margin = Inches(1)

# ── Style helpers ─────────────────────────────────────────────────────────────
def set_font(run, name="Times New Roman", size=12, bold=False, italic=False, color=None):
    run.font.name  = name
    run.font.size  = Pt(size)
    run.font.bold  = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def body_para(text, bold=False, italic=False, space_after=6, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = Pt(22)
    p.alignment = alignment
    if text:
        run = p.add_run(text)
        set_font(run, bold=bold, italic=italic)
    return p

def heading1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    set_font(run, size=13, bold=True)
    return p

def heading2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    set_font(run, size=12, bold=True)
    return p

def heading3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(text)
    set_font(run, size=12, bold=False, italic=True)
    return p

def add_rule():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def table_style(tbl, header_row=True):
    tbl.style = 'Table Grid'
    for i, row in enumerate(tbl.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for para in cell.paragraphs:
                para.paragraph_format.space_before = Pt(3)
                para.paragraph_format.space_after  = Pt(3)
                para.paragraph_format.line_spacing = Pt(14)
                for run in para.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(10)
                    if i == 0 and header_row:
                        run.font.bold = True
    # shade header row
    if header_row and tbl.rows:
        for cell in tbl.rows[0].cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), 'D9E1F2')
            tcPr.append(shd)

def add_table(headers, rows, col_widths=None):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    # header row
    for j, h in enumerate(headers):
        cell = tbl.rows[0].cells[j]
        cell.text = h
        if col_widths:
            cell.width = Inches(col_widths[j])
    # data rows
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = tbl.rows[i+1].cells[j]
            cell.text = str(val)
    table_style(tbl)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    return tbl

def code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.4)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(10)
    return p

def inline_bold(para, text):
    run = para.add_run(text)
    set_font(run, bold=True)
    return run

def inline_italic(para, text):
    run = para.add_run(text)
    set_font(run, italic=True)
    return run

# ═══════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ═══════════════════════════════════════════════════════════════════════════

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(24)
p.paragraph_format.space_after  = Pt(16)
run = p.add_run(
    "Context-Specific Innate Immune Evasion via VDAC1 Gate-Jamming "
    "in Microsatellite-Stable Colorectal Cancer: "
    "A Three-Cohort Transcriptomic Analysis"
)
set_font(run, size=16, bold=True)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.paragraph_format.space_after = Pt(4)
run2 = p2.add_run("Anthony J. Vasquez Sr.")
set_font(run2, size=12, bold=True)

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
p3.paragraph_format.space_after = Pt(4)
run3 = p3.add_run("Delaware Valley University, Doylestown, PA, USA")
set_font(run3, size=11, italic=True)

p4 = doc.add_paragraph()
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
p4.paragraph_format.space_after = Pt(16)
run4 = p4.add_run("Corresponding author: vasquezaj3921@delval.edu")
set_font(run4, size=11)

add_rule()

# ═══════════════════════════════════════════════════════════════════════════
# ABSTRACT
# ═══════════════════════════════════════════════════════════════════════════

heading1("Abstract")

body_para(
    "Immune checkpoint inhibitors (ICIs) have transformed oncology for microsatellite "
    "instability-high (MSI-H) colorectal cancer, yet 85-95% of colorectal cancer patients "
    "carry microsatellite-stable (MSS) tumors and derive no benefit from current ICI regimens. "
    "We propose that VDAC1-mediated mitochondrial DNA (mtDNA) gate-jamming -- suppression of "
    "VDAC1 oligomerization by HK-II docking, Bcl-xL binding, and outer mitochondrial membrane "
    "cholesterol loading -- explains this selectivity by silencing the cGAS-STING innate immune "
    "signal required for spontaneous T cell priming. To test this hypothesis at scale, we computed "
    "a transcriptomic Gate-Jamming Score (tGJS = 0.4 x HK2 + 0.3 x BCL2L1 + 0.3 x TSPO, "
    "rank-normalized) and conducted three sequential analyses: (S1) pan-cancer TCGA "
    "(n = 10,071, 33 cancer types) -- null result (rho = +0.38 vs ICI response rate, p = 0.14); "
    "(S2) COADREAD MSS/TP53-wildtype clean room (n = 209) -- five Bonferroni-significant inverse "
    "correlations between tGJS and immune markers including HAVCR2 (rho = -0.349, p_bonf = 5 x 10^-6), "
    "CXCL10 (rho = -0.231, p_bonf = 0.015), and cGAS (rho = -0.208, p_bonf = 0.049); "
    "(S3) IMvigor210 urothelial carcinoma atezolizumab cohort (n = 348) -- null result "
    "(Wilcoxon p = 0.965, Cox HR = 0.898, p = 0.455). The flanking nulls (S1, S3) define the "
    "framework's domain: the gate-jamming signal is detectable only when VDAC1-mediated mtDNA "
    "release is the dominant cytosolic DNA source and innate priming is the rate-limiting step. "
    "The S2 clean room results, combined with the three-layer therapeutic hypothesis "
    "(VDAC1 gate-opener + cGAMP/DNA eraser inhibitor + checkpoint blockade) independently derived "
    "from the same data by three AI analytical systems, motivates protein-level validation in MSS "
    "colorectal cancer and combination ICI trials in this specific population."
)

p_kw = doc.add_paragraph()
p_kw.paragraph_format.space_after = Pt(6)
inline_bold(p_kw, "Keywords: ")
p_kw.add_run("VDAC1, gate-jamming, cGAS-STING, microsatellite-stable colorectal cancer, innate immunity, immunotherapy resistance, tGJS")
for run in p_kw.runs[1:]:
    set_font(run)

add_rule()

# ═══════════════════════════════════════════════════════════════════════════
# 1. INTRODUCTION
# ═══════════════════════════════════════════════════════════════════════════

heading1("1. Introduction")

body_para(
    "Colorectal cancer is the second leading cause of cancer mortality in the United States, "
    "with approximately 150,000 new cases annually. Pembrolizumab and nivolumab achieve durable "
    "responses in the ~5-15% of patients with MSI-H tumors -- where pervasive genomic instability "
    "generates abundant mutational neoantigens and activates innate immune sensing. For the "
    "remaining 85-95% of patients with MSS tumors, ICI monotherapy has consistently failed in "
    "randomized trials, and no predictive biomarker identifies a responsive MSS subpopulation. "
    "The central unmet need in colorectal cancer immunotherapy is converting MSS tumors from "
    "immune-cold to immune-hot."
)

body_para(
    "VDAC1 (voltage-dependent anion channel 1) is the most abundant protein in the outer "
    "mitochondrial membrane, present at densities exceeding 1,000 molecules per um2. Its "
    "oligomeric form releases 500-650 bp mitochondrial DNA (mtDNA) fragments into the cytoplasm "
    "(Kim et al. 2019, Science), activating the cGAS-STING innate immune pathway required for "
    "spontaneous CD8+ T cell priming (Woo et al. 2014, Immunity). Cancer cells suppress this "
    "oligomerization through at least three mechanisms: hexokinase-II (HK-II) docking on VDAC1's "
    "outer barrel (Wolf et al. 2023, Science Immunology; Bieker et al. 2025, Communications "
    "Biology), Bcl-xL binding via its BH4 domain (Monaco et al. 2015, JBC), and cholesterol "
    "loading of the outer mitochondrial membrane (Betaneli et al. 2012, Biophysical Journal). "
    "Together, these constitute a gate-jamming architecture that silences the mitochondrial innate "
    "immune alarm."
)

body_para(
    "Critically, this mechanism is predicted to be relevant precisely in MSS tumors. In MSI-H "
    "tumors, nuclear DNA damage generates cytosolic DNA fragments independently of mitochondrial "
    "dynamics, saturating cGAS-STING regardless of VDAC1 state. In MSS tumors where nuclear DNA "
    "damage is minimal, VDAC1-mediated mtDNA release is the predicted dominant -- potentially "
    "sole -- cytosolic DNA source. Gate-jamming in this context suppresses the entire innate "
    "priming axis, explaining ICI failure without invoking T cell exhaustion."
)

body_para(
    "We formalized this hypothesis as a transcriptomic Gate-Jamming Score (tGJS) and conducted "
    "three sequential analyses designed to test the framework, define its domain, and identify "
    "the biological context where protein-level validation is most warranted."
)

add_rule()

# ═══════════════════════════════════════════════════════════════════════════
# 2. METHODS
# ═══════════════════════════════════════════════════════════════════════════

heading1("2. Methods")

heading2("2.1 Transcriptomic Gate-Jamming Score")

body_para(
    "The tGJS is a rank-normalized composite of three genes encoding the primary VDAC1 "
    "gate-jamming proteins:"
)
code_block("tGJS = 0.40 x norm(HK2) + 0.30 x norm(BCL2L1) + 0.30 x norm(TSPO)")
body_para(
    "HK2 encodes hexokinase-II (VDAC1 docking, weight 0.40); BCL2L1 encodes Bcl-xL (VDAC1 "
    "binding, weight 0.30); TSPO encodes the Translocator Protein, a cholesterol transport "
    "protein used as a proxy for mitochondrial cholesterol loading (weight 0.30). Normalization "
    "is performed within each cohort."
)

heading2("2.2 S1: Pan-Cancer TCGA Analysis")
body_para(
    "Expression data for 10,071 samples across 33 TCGA cancer types were retrieved from the "
    "TCGA PanCanAtlas (cBioPortal, pan_cancer_atlas_2018). ICI response rates per cancer type "
    "were compiled from published meta-analyses. tGJS was correlated with per-cancer-type mean "
    "ICI response rate using Spearman rank correlation."
)

heading2("2.3 S2: COADREAD MSS-Stratified Clean Room Analysis")
body_para(
    "592 COADREAD samples from coadread_tcga_pan_can_atlas_2018 were stratified by MSI status "
    "x TP53 mutation into four groups. MSI status was determined from MANTIS (threshold >= 0.4) "
    "and MSISensor (threshold > 10) scores; MSI-H if either score exceeded threshold. TP53 "
    "mutation status was retrieved from the COADREAD mutation profile. Spearman correlations "
    "were computed between tGJS and 20 immune/pathway markers in each stratum, with Bonferroni "
    "correction for 20 comparisons per stratum."
)

heading2("2.4 S3: IMvigor210 Atezolizumab Cohort")
body_para(
    "The IMvigor210CoreBiologies R package (Mariathasan et al. 2018) provides pre-treatment "
    "bulk RNA-Seq from 348 urothelial carcinoma patients treated with atezolizumab, with RECIST "
    "response, overall survival, TMB, IC Level, and immune phenotype classification. Raw counts "
    "were converted to TPM. tGJS was computed using z-score normalization within the cohort. "
    "Primary tests: Wilcoxon rank-sum, logistic regression, Spearman rho vs ordinal RECIST, "
    "Cox proportional hazards, and log-rank (tGJS high vs low)."
)

heading2("2.5 Multi-Model Convergence Protocol")
body_para(
    "The gate-jamming therapeutic hypothesis was evaluated using the IRIS (Independent "
    "Replicated Inquiry System) protocol, in which five independent large language models "
    "(Claude Opus 4.6, Gemini 2.5 Pro, Grok 4.1 Fast, Mistral Large, DeepSeek Chat) received "
    "the identical compiled question without cross-exposure. Claims were extracted and embedded "
    "using all-MiniLM-L6-v2 (384-dimensional), clustered by cosine similarity >= 0.70, and "
    "classified by model agreement (TYPE 0 = 4-5 models, TYPE 1 = 3/5, TYPE 2 = 2/5, "
    "TYPE 3 = 1/5). Cross-run analysis across 28 independent runs (27,931 pairwise comparisons) "
    "assessed inter-run replication."
)

add_rule()

# ═══════════════════════════════════════════════════════════════════════════
# 3. RESULTS
# ═══════════════════════════════════════════════════════════════════════════

heading1("3. Results")

heading2("3.1 S1: Pan-Cancer tGJS Does Not Predict ICI Response (n = 10,071)")
body_para(
    "Across 33 TCGA cancer types, tGJS was not inversely correlated with published ICI response "
    "rates (Spearman rho = +0.382, p = 0.144). The positive trend reflects that metabolically "
    "aggressive tumors -- which have higher tGJS -- tend to be the same tumors with higher "
    "mutational burden and baseline immune activation, creating a confound that overwhelms any "
    "gate-jamming signal at the cross-cancer level. ENPP1 expression showed the strongest "
    "pan-cancer inverse correlation with tGJS (rho = -0.181, p = 4.3 x 10^-75), initially "
    "interpreted as evidence of orthogonal evasion strategies. The pan-cancer null is consistent "
    "with the hypothesis that a cancer-type-homogeneous, immunogenomically stratified analysis "
    "is required to observe the gate-jamming signal. Full results are in Supplementary Analysis S1."
)

heading2("3.2 S2: MSS/TP53-wt Clean Room Recovers Five Bonferroni-Significant Signals (n = 209)")
body_para(
    "Restricting analysis to the MSS/TP53-wildtype COADREAD stratum recovered five markers "
    "reaching Bonferroni significance (Table 1), all in directions predicted by the gate-jamming "
    "hypothesis."
)

# Table 1
p_t1 = doc.add_paragraph()
p_t1.paragraph_format.space_before = Pt(8)
p_t1.paragraph_format.space_after = Pt(4)
inline_bold(p_t1, "Table 1.")
p_t1.add_run(" Bonferroni-significant correlations in the MSS/TP53-wt clean room (n = 209).")
for run in p_t1.runs[1:]:
    set_font(run, italic=True)

add_table(
    headers=["Marker", "Spearman rho", "p_bonf", "Direction", "Interpretation"],
    rows=[
        ["HAVCR2 (TIM-3)", "-0.349", "5 x 10^-6", "Down", "Fewer T cells infiltrate to become exhausted"],
        ["TREX1",          "+0.315", "7 x 10^-5", "Up",   "Co-deployment of mtDNA erasure with gate-jamming"],
        ["CXCL10",         "-0.231", "0.015",      "Down", "Reduced IFN-gamma-induced chemokine recruitment"],
        ["STING ratio",    "-0.216", "0.034",      "Down", "Residual STING signaling shifts to immunosuppressive profile"],
        ["cGAS (MB21D1)",  "-0.208", "0.049",      "Down", "Lower cGAS in high-tGJS tumors"],
    ],
    col_widths=[1.4, 1.1, 0.9, 0.7, 2.9]
)

body_para(
    "No immune markers reached Bonferroni significance in the MSI-H control strata (n = 67 "
    "and n = 28), consistent with the prediction that genomic instability saturates cGAS-STING "
    "independently of VDAC1 state."
)

body_para(
    "The ENPP1 anti-correlation (rho = -0.027, not significant) did not replicate within the "
    "clean room, correcting the S1 interpretation: the pan-cancer ENPP1 signal was driven by "
    "cross-cancer-type expression differences rather than a within-tumor-type biological "
    "relationship."
)

body_para(
    "Two unexpected findings refined the evasion architecture model. First, the positive "
    "correlation between tGJS and TREX1 (rho = +0.315) indicates that the most evasion-committed "
    "MSS tumors co-deploy mitochondrial gate-jamming and cytosolic DNA erasure simultaneously -- "
    "a belt-and-suspenders strategy rather than an either/or trade-off. Second, HAVCR2 (TIM-3) "
    "showed the strongest anti-correlation, suggesting high-tGJS MSS tumors suffer from T cell "
    "absence rather than T cell exhaustion -- consistent with suppressed innate priming preventing "
    "initial T cell recruitment. Full stratified results across all 20 markers and 4 strata are "
    "in Supplementary Analysis S2."
)

heading2("3.3 S3: tGJS Does Not Predict Atezolizumab Response in Urothelial Carcinoma (n = 348)")
body_para(
    "In the IMvigor210 cohort, tGJS showed no association with atezolizumab response or overall "
    "survival at any level of analysis (Table 2)."
)

p_t2 = doc.add_paragraph()
p_t2.paragraph_format.space_before = Pt(8)
p_t2.paragraph_format.space_after = Pt(4)
inline_bold(p_t2, "Table 2.")
p_t2.add_run(" IMvigor210 primary results (n = 348, n_response = 298).")
for run in p_t2.runs[1:]:
    set_font(run, italic=True)

add_table(
    headers=["Test", "Result", "p-value"],
    rows=[
        ["Wilcoxon (CR/PR vs SD/PD)",    "--",                              "p = 0.965"],
        ["Logistic regression",           "OR = 1.038",                     "p = 0.868"],
        ["Spearman rho vs RECIST",        "rho = -0.017",                   "p = 0.767"],
        ["Kruskal-Wallis by tertile",     "--",                              "p = 0.559"],
        ["Cox PH (continuous)",           "HR = 0.898 (95% CI: 0.678-1.190)", "p = 0.455"],
        ["Log-rank (high vs low)",        "--",                              "p = 0.587"],
        ["Median OS, tGJS-High",          "20.6 months",                    "--"],
        ["Median OS, tGJS-Low",           "20.6 months",                    "--"],
    ],
    col_widths=[2.8, 2.8, 1.4]
)

body_para(
    "This null is mechanistically expected: urothelial carcinoma carries high baseline somatic "
    "mutation burden from tobacco and occupational carcinogen exposure, generating nuclear "
    "DNA-derived cytosolic DNA independently of VDAC1 state. Additionally, atezolizumab targets "
    "the adaptive immune checkpoint on exhausted T cells downstream of innate priming; even if "
    "gate-jamming were relevant, its effect would operate upstream of the treatment's mechanism "
    "of action. Full results and four figures are in Supplementary Analysis S3."
)

add_rule()

# ═══════════════════════════════════════════════════════════════════════════
# 4. DISCUSSION
# ═══════════════════════════════════════════════════════════════════════════

heading1("4. Discussion")

heading2("4.1 Three Analyses Define the Framework's Domain")

body_para(
    "The S1 -> S2 -> S3 arc is the central finding of this work. It answers not only "
    "'is there a signal?' but 'where does the signal live?':"
)

p_t3 = doc.add_paragraph()
p_t3.paragraph_format.space_before = Pt(8)
p_t3.paragraph_format.space_after = Pt(4)
inline_bold(p_t3, "Table 3.")
p_t3.add_run(" Three-cohort analysis arc.")
for run in p_t3.runs[1:]:
    set_font(run, italic=True)

add_table(
    headers=["Analysis", "Cohort", "Context", "tGJS Signal"],
    rows=[
        ["S1", "TCGA pan-cancer (n = 10,071)", "33 cancer types, ICI response proxy", "Null -- cross-cancer confounds"],
        ["S2", "TCGA COADREAD MSS/TP53-wt (n = 209)", "Clean room: MSS + intact apoptosis", "5 Bonferroni-significant immune correlations"],
        ["S3", "IMvigor210 urothelial (n = 348)", "High TMB, PD-L1 blockade", "Null -- wrong context, wrong treatment mechanism"],
    ],
    col_widths=[0.5, 2.3, 2.2, 2.0]
)

body_para(
    "The two nulls are not failures -- they are the framework's falsification boundaries. "
    "A signal that appears everywhere is not a mechanism; it is noise. A signal that appears "
    "precisely where the biology predicts it should appear, and not where the biology predicts "
    "it shouldn't, is evidence that the model is capturing something real. The COADREAD "
    "MSS/TP53-wildtype stratum represents the tightest available approximation of the predicted "
    "clean room using existing public data: minimal nuclear DNA instability, intact TP53 "
    "apoptosis signaling, VDAC1 as the predicted dominant cytosolic DNA source."
)

heading2("4.2 The Therapeutic Hypothesis: Three Independent Systems, One Stack")

body_para(
    "The analytical arc motivates a specific three-layer therapeutic intervention. This "
    "hypothesis was independently derived by three AI analytical systems (Claude Opus, "
    "Gemini Pro, Grok) working from the same data without cross-exposure, converging on "
    "identical components:"
)

items = [
    ("1. VDAC1 gate-opener: ", "displace HK-II from VDAC1 (methyl jasmonate, 2-DG, clotrimazole) to restore oligomerization-dependent mtDNA release and cGAS-STING activation."),
    ("2. DNA/cGAMP eraser inhibitor: ", "inhibit TREX1 or ENPP1 to prevent degradation of the released mtDNA or downstream cGAMP, amplifying and sustaining innate immune activation."),
    ("3. Checkpoint blockade: ", "administer anti-PD-1/PD-L1 to prevent exhaustion of T cells now being recruited by the restored innate signal."),
]
for label, rest in items:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(4)
    inline_bold(p, label)
    run = p.add_run(rest)
    set_font(run)

body_para(
    "The order matters: gate-opener generates the innate signal; eraser inhibitor sustains it; "
    "checkpoint blockade amplifies the adaptive response. The TREX1 co-occurrence finding from "
    "S2 directly informs step two: high-gate-jamming tumors simultaneously upregulate cytosolic "
    "DNA erasure, implying gate-opener alone may be insufficient. Three independently reasoning "
    "systems arriving at the same three-layer stack from the same data constitutes a prior strong "
    "enough to justify the next experiment."
)

heading2("4.3 Why MSS Colorectal Is the Right Target")

body_para(
    "MSS colorectal adenocarcinoma has: (1) the largest absolute population of ICI-refractory "
    "patients among common cancers where ICI is attempted; (2) well-characterized MSI "
    "stratification, allowing clean separation of the signal; (3) an available TP53 mutation "
    "axis that further refines the predicted clean room; and (4) a tumor biology where VDAC1 "
    "docking by HK-II is supported by multiple independent lines of evidence. The HAVCR2 "
    "finding (rho = -0.349) points to T cell absence rather than T cell exhaustion as the "
    "immune bottleneck in high-tGJS MSS tumors -- a meaningful clinical distinction suggesting "
    "anti-TIM-3 trials in MSS CRC are addressing the wrong checkpoint."
)

heading2("4.4 The GJS as One Layer of a Multi-Layer Evasion Architecture")

body_para(
    "The GJS measures one specific bottleneck -- VDAC1 oligomerization suppression -- within "
    "a multi-layer immune evasion system. The cGAS-STING axis is not uniformly anti-tumor: "
    "Lai et al. (2025, Immunity) showed that VDAC-mediated mtDNA from senescent tumor cells "
    "can enhance immunosuppression through MDSC recruitment. The tumor microenvironment may "
    "override gate-restoration even when cGAS-STING is successfully activated. STING pathway "
    "competence must be assessed -- STING silencing via promoter methylation occurs in 1-25% "
    "of tumors pan-cancer, and high-GJS tumors with silenced STING require epigenetic "
    "reactivation before gate-restoration can be effective."
)

p_t4 = doc.add_paragraph()
p_t4.paragraph_format.space_before = Pt(8)
p_t4.paragraph_format.space_after = Pt(4)
inline_bold(p_t4, "Table 4.")
p_t4.add_run(" GJS x STING status therapeutic matrix.")
for run in p_t4.runs[1:]:
    set_font(run, italic=True)

add_table(
    headers=["", "STING Intact", "STING Silenced"],
    rows=[
        ["High GJS", "Primary target -- gate-opener + eraser inhibitor + checkpoint", "Requires DNMT inhibitor before gate-restoration"],
        ["Low GJS",  "Gate open, pathway active -- checkpoint alone may suffice",       "Chronic signaling, paradoxical immunosuppression"],
    ],
    col_widths=[1.0, 3.2, 2.8]
)

heading2("4.5 What the tGJS Does Not Capture")

body_para(
    "The transcriptomic proxy measures mRNA abundance of three genes; it cannot distinguish "
    "HK-II docked on VDAC1 versus in the cytosol, Bcl-xL bound to VDAC1 versus other targets, "
    "or TSPO-mediated cholesterol at the outer mitochondrial membrane versus elsewhere. "
    "Protein-level measurement -- proximity ligation assay for HK-II-VDAC1 and Bcl-xL-VDAC1 "
    "complexes, mitochondrial lipidomics for the Chol/CL ratio -- is required to compute the "
    "true GJS:"
)
code_block("GJS = f_HKII x 0.40 + f_BclxL x 0.30 + [Chol]/[CL]_norm x 0.30")
body_para(
    "where f_HKII is the fraction of VDAC1 occupied by HK-II and f_BclxL the fraction bound "
    "by Bcl-xL (both 0-1). The tGJS is a screening tool to identify where this protein-level "
    "assay should be run first."
)

heading2("4.6 Limitations")

limits = [
    "No experimental validation. All findings are computational. The weights (0.4/0.3/0.3) are convergence-derived estimates, not empirically optimized coefficients.",
    "Bulk RNA-Seq. Expression values are cellular averages across tumor, stromal, and immune cells. Single-cell attribution is required to determine which cell type drives the tGJS signal.",
    "TCGA lacks treatment data. The S2 correlations are with immune markers, not ICI outcomes. The hypothesis that high-tGJS MSS CRC tumors fail ICI is mechanistically motivated but not directly tested.",
    "IMvigor210 is a single cohort. The S3 null applies to urothelial carcinoma treated with PD-L1 blockade. It does not predict the outcome in MSS CRC treated with combination gate-restoration regimens.",
    "VBIT-4 specificity. Ravishankar et al. (2025, bioRxiv) showed VBIT-4 disrupts membranes independent of VDAC1 at >= 30 uM. Claims requiring VBIT-4 need orthogonal genetic validation.",
    "TREX1 inhibitors are not clinically available. This identifies a drug development gap.",
]
for i, lim in enumerate(limits, 1):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(4)
    inline_bold(p, f"{i}. ")
    run = p.add_run(lim)
    set_font(run)

heading2("4.7 Immediate Next Steps")

steps = [
    ("In vitro (4-8 weeks): ", "Displace HK-II from VDAC1 in immune-cold cell lines (Panc-1, HCT116) and measure cytoplasmic mtDNA, p-STING (Ser366), and IFN-beta. Compute GJS across 15+ cell lines and correlate with basal cGAS-STING activity (Spearman rho <= -0.6 predicted)."),
    ("Protein-level validation in COADREAD tissue (8-12 weeks): ", "Run proximity ligation assay for HK-II-VDAC1 complexes in MSS vs MSI-H colorectal cancer tissue sections. Correlate PLA signal with CD8A IHC and tGJS."),
    ("In vivo (10 weeks): ", "Gate-restoration combination (methyl jasmonate + ABT-263) + anti-PD-1 in 4T1 (immune-cold) and MC38 (immune-hot) syngeneic tumor models. Prediction: synergy in 4T1, no added benefit in MC38."),
    ("GSE91061 melanoma validation: ", "The Riaz 2017 nivolumab cohort (n = 109) is the next planned validation to confirm the boundary conditions in a second high-TMB tumor type."),
]
for label, rest in steps:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(4)
    inline_bold(p, label)
    run = p.add_run(rest)
    set_font(run)

add_rule()

# ═══════════════════════════════════════════════════════════════════════════
# 5. DATA AVAILABILITY
# ═══════════════════════════════════════════════════════════════════════════

heading1("5. Data Availability")

urls = [
    ("Analysis scripts: ", "analysis/tcga_gjs/compute_tgjs.py, compute_tgjs_coadread_mss.py; analysis/imvigor210/compute_tgjs_imvigor210.R"),
    ("Repository: ", "https://github.com/templetwo/vdac-pharmacology-atlas"),
    ("Dataset archive: ", "https://huggingface.co/datasets/TheTempleofTwo/vdac-pharmacology-atlas"),
    ("OSF preregistration: ", "https://osf.io/c9rqb/"),
    ("IRIS Gate Evo pipeline: ", "https://github.com/templetwo/iris-gate-evo"),
]
for label, val in urls:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(4)
    inline_bold(p, label)
    run = p.add_run(val)
    set_font(run)

add_rule()

# ═══════════════════════════════════════════════════════════════════════════
# 6. AUTHOR CONTRIBUTIONS
# ═══════════════════════════════════════════════════════════════════════════

heading1("6. Author Contributions")

body_para(
    "A.J.V. conceived the research questions, designed the analytical framework, executed all "
    "computational analyses, interpreted all results, and wrote the manuscript. Computational "
    "analyses were executed with Claude Code (Anthropic). Manuscript drafting was assisted by "
    "Claude (Anthropic). The IRIS multi-model convergence protocol was developed and run by "
    "A.J.V. using five independent AI systems (Claude Opus, Gemini Pro, Grok, Mistral, "
    "DeepSeek); convergence metrics are system-computed. All scientific decisions, "
    "interpretations, and conclusions are solely the responsibility of A.J.V."
)

heading1("7. Competing Interests")
body_para("The author declares no competing interests. This work received no external funding.")

heading1("8. License")
body_para("This manuscript and all associated data are released under CC BY 4.0.")

add_rule()

# ═══════════════════════════════════════════════════════════════════════════
# REFERENCES
# ═══════════════════════════════════════════════════════════════════════════

heading1("References")

refs = [
    "Betaneli V, Petrov EP, Schwille P. (2012) The role of lipids in VDAC oligomerization. Biophys J 102:523.",
    "Bieker JT, Timme S, et al. (2025) A membrane-buried glutamate mediates VDAC-hexokinase binding. Commun Biol 8:212.",
    "Carozza JA, et al. (2023) ENPP1 as an innate immune checkpoint. PNAS.",
    "Daniilidis M, Gunsel U, et al. (2025) Structural basis of apoptosis induction by VDAC1. Nat Commun 16:9481.",
    "Fadzeyeva E, et al. (2026) CBD-induced VDAC1 oligomerization-dependent effects. Pharmaceuticals 19:95.",
    "Gehrcken L, et al. (2025) cGAS-STING in anti-tumor immunity. Adv Sci 12:2500296.",
    "Goicoechea L, et al. (2023) Mitochondrial cholesterol: metabolism and impact on redox biology. Redox Biol 61:102643.",
    "Ikeda H, et al. (2025) Mitochondrial transfer from cancer cells to TILs. Nature.",
    "Jahn H, Bartos L, Dearden GD, et al. (2023) VDAC1 dimers as phospholipid scramblases. Nat Commun 14:8115.",
    "Kim J, et al. (2019) VDAC oligomers form mitochondrial pores to release mtDNA. Science 366:1531.",
    "Lafargue K, et al. (2025) Lipid regulation of VDAC1 assemblies. Commun Biol 8:936.",
    "Lai J, et al. (2025) VDAC-mediated mtDNA release from senescent tumor cells. Immunity 58:811.",
    "Mangalhara KC, et al. (2023) Complex II loss activates MHC-I. Science 381:1316.",
    "Mariathasan S, et al. (2018) TGFbeta attenuates tumour response to PD-L1 blockade by contributing to exclusion of T cells. Nature 554:544.",
    "Monaco G, et al. (2015) The BH4 domain of Bcl-xL targets VDAC1 for Ca2+ control. J Biol Chem 290:9150.",
    "Montero J, et al. (2008) Mitochondrial cholesterol contributes to chemotherapy resistance in HCC. Cancer Res 68:5246.",
    "Prashar A, et al. (2024) VDAC1-dependent inner membrane herniation vesicles. Nature 632:1110.",
    "Ravishankar H, et al. (2025) VBIT-4 specificity challenge. bioRxiv.",
    "Ren H, et al. (2025) VSTM2L enhances HK2-VDAC1 interaction. Nat Commun 16:1534.",
    "Rimmerman N, et al. (2013) Direct modulation of the OMM channel VDAC1 by CBD. Cell Death Dis 4:e949.",
    "Samson N, Ablasser A. (2022) The cGAS-STING pathway and cancer. Nat Cancer 3:1452.",
    "Shoshan-Barmatz V, et al. (2025) p53 directly binds VDAC1. Biomolecules 16:141.",
    "Wan X, et al. (2026) VDAC1 oligomerization regulates PANoptosis. Neural Regen Res 21(4).",
    "Wolf AJ, et al. (2023) HK2 dissociation from VDAC1 triggers NLRP3. Sci Immunol 8:eade7652.",
    "Woo SR, et al. (2014) STING-dependent cytosolic DNA sensing mediates innate immune recognition of tumors. Immunity 41:830.",
    "Xian H, et al. (2022) Oxidized mtDNA exit via mPTP and VDAC channels. Immunity 55:1370.",
]

for ref in refs:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent   = Inches(0.4)
    p.paragraph_format.first_line_indent = Inches(-0.4)
    p.paragraph_format.space_after   = Pt(4)
    p.paragraph_format.line_spacing  = Pt(16)
    run = p.add_run(ref)
    set_font(run, size=11)

# ── Save ─────────────────────────────────────────────────────────────────────
out = "/Users/vaquez/vdac-pharmacology-atlas/paper/gjs_manuscript.docx"
doc.save(out)
print(f"Saved: {out}")
